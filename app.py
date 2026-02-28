import os
import csv
import io
import re
import threading
import json
from datetime import datetime
from urllib.parse import quote  # <--- ВАЖЛИВИЙ ІМПОРТ ДЛЯ ВИПРАВЛЕННЯ ПОМИЛКИ
<<<<<<< HEAD
import logging
import os
from dotenv import load_dotenv

from flask import Flask, render_template, request, redirect, url_for, session, flash, Response, stream_with_context
from flask_session import Session
from filelock import FileLock
=======

from flask import Flask, render_template, request, redirect, url_for, session, flash, Response, stream_with_context
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620
from data import get_indicator_choices, get_indicator_details

# Імпорти для Word
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

<<<<<<< HEAD
load_dotenv()

app = Flask(__name__)
# Секретний ключ для сесій
app.secret_key = os.environ.get('SECRET_KEY', 'd5aeb79aff27c1cbc690473e25c5b70dbcc959da288a0f67')

# Configure Flask-Session
app.config['SESSION_TYPE'] = 'filesystem'
app.config['SESSION_FILE_DIR'] = './flask_session'
Session(app)

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

LEADERBOARD_FILE = 'leaderboard.csv'
LEADERBOARD_LOCK_FILE = 'leaderboard.csv.lock'
ALLOWED_EXTENSIONS = {'csv'}

# Пароль адміністратора для видалення записів
ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'admin')

def escape_csv_injection(val):
    s = str(val)
    if s and s.startswith(('=', '+', '-', '@')):
        return f"'{s}"
    return s

def calculate_score(indicator_type, base_weight, n=1.0, s=1.0, k=1.0):
    score = 0
    if indicator_type in ('fixed_value', 'fixed'):
        score = base_weight
    elif indicator_type == 'simple_multiplication':
        score = base_weight * n
    elif indicator_type in ('percentage_update', 'positive_feedback'):
        percent = s if s <= 1.0 else s / 100.0
        score = base_weight * n * percent
    elif indicator_type == 'scientific_publication':
        score = base_weight * n * k * s
    elif indicator_type == 'quantity_share':
        score = base_weight * n * s
    return score
=======
app = Flask(__name__)
# Секретний ключ для сесій
app.secret_key = 'd5aeb79aff27c1cbc690473e25c5b70dbcc959da288a0f67'

LEADERBOARD_FILE = 'leaderboard.csv'
ALLOWED_EXTENSIONS = {'csv'}
leaderboard_lock = threading.Lock()

# Пароль адміністратора для видалення записів
ADMIN_PASSWORD = "admin" 
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def load_leaderboard():
    leaderboard_data = {}
<<<<<<< HEAD
    lock = FileLock(LEADERBOARD_LOCK_FILE)
    try:
        with lock:
            if not os.path.exists(LEADERBOARD_FILE):
                return {}
=======
    try:
        with leaderboard_lock:
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620
            with open(LEADERBOARD_FILE, 'r', newline='', encoding='utf-8-sig') as csvfile:
                reader = csv.reader(csvfile, delimiter=';')
                header = next(reader, None)
                for i, row in enumerate(reader):
                    try:
                        if len(row) >= 2: 
                            name = row[0].strip()
<<<<<<< HEAD
                            # remove the leading quote if it was escaped
                            name = name[1:] if name.startswith("'") and len(name) > 1 and name[1] in ('=', '+', '-', '@') else name
                            position = row[1].strip() if len(row) > 2 else "Не вказано"
                            position = position[1:] if position.startswith("'") and len(position) > 1 and position[1] in ('=', '+', '-', '@') else position
                            
=======
                            position = row[1].strip() if len(row) > 2 else "Не вказано"
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620
                            score_str = row[2] if len(row) > 2 else row[1]
                            score = float(score_str.replace(',', '.'))
                            if name:
                                if name not in leaderboard_data or score > leaderboard_data[name]['score']:
                                    leaderboard_data[name] = {'score': score, 'position': position}
                    except Exception: pass
<<<<<<< HEAD
    except Exception as e: 
        logging.error(f"Error loading leaderboard: {e}")
    return leaderboard_data

def save_leaderboard(leaderboard_data):
    lock = FileLock(LEADERBOARD_LOCK_FILE)
    try:
        with lock:
=======
    except FileNotFoundError: pass
    return leaderboard_data

def save_leaderboard(leaderboard_data):
    try:
        with leaderboard_lock:
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620
            with open(LEADERBOARD_FILE, 'w', newline='', encoding='utf-8-sig') as csvfile:
                writer = csv.writer(csvfile, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
                writer.writerow(['ПІБ', 'Посада', 'Загальний бал'])
                for name, data in sorted(leaderboard_data.items()):
                    score_str = "{:.2f}".format(data.get('score', 0)).replace('.', ',')
<<<<<<< HEAD
                    writer.writerow([escape_csv_injection(name), escape_csv_injection(data.get('position', 'Не вказано')), score_str])
    except Exception as e: 
        logging.error(f"Error saving leaderboard: {e}")
=======
                    writer.writerow([name, data.get('position', 'Не вказано'), score_str])
    except Exception as e: print(f"Error saving: {e}")
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620

def parse_input_data_string(input_str, expected_type):
    parsed_values = {}
    input_str = input_str.strip()
    if expected_type == 'boolean': 
        parsed_values['boolean_value'] = (input_str.lower() == 'так')
        return parsed_values
    elif expected_type == 'fixed' or input_str == '-': return {}
    
    n_match = re.search(r'n=(\d+([.,]\d+)?)', input_str, re.IGNORECASE)
    s_match = re.search(r'S=(\d+([.,]\d+)?)', input_str, re.IGNORECASE)
    k_match = re.search(r'k=(\d+([.,]\d+)?)', input_str, re.IGNORECASE)
    
    try:
        if n_match: parsed_values['n_value'] = float(n_match.group(1).replace(',', '.'))
        if s_match: parsed_values['s_value'] = float(s_match.group(1).replace(',', '.'))
        if k_match: parsed_values['k_value'] = float(k_match.group(1).replace(',', '.'))
    except ValueError: pass
    return parsed_values

@app.route('/')
def index():
    indicator_choices = get_indicator_choices()
    current_entries = session.get('entries', [])
    current_total_block1 = sum(e['score'] for e in current_entries if e.get('block') == 1)
    current_total_block2 = sum(e['score'] for e in current_entries if e.get('block') == 2)
    return render_template('index.html', indicator_choices=indicator_choices, current_entries=current_entries, current_total_block1=current_total_block1, current_total_block2=current_total_block2)

@app.route('/update_personal_data', methods=['POST'])
def update_personal_data():
    session['full_name'] = request.form.get('full_name', '').strip()
    session['institution_type'] = request.form.get('institution_type')
    session['department'] = request.form.get('department')
    session['position'] = request.form.get('position')
    session.modified = True
    flash('Персональні дані збережено.', 'success')
    return redirect(url_for('index'))


@app.route('/add', methods=['POST'])
def add_entry():
    indicator_id = request.form.get('indicator_id')
    if not indicator_id: return redirect(url_for('index'))

    indicator_details = get_indicator_details(indicator_id)
    if not indicator_details: return redirect(url_for('index'))

    indicator_type = indicator_details.get('formula_type', 'fixed_value')
    base_weight = indicator_details.get('weight', 0)

    score = 0
    input_values = {}
    comment = request.form.get('comment', '').strip()

    try:
        n_val = request.form.get('n_value')
        s_val = request.form.get('s_value')
        k_val = request.form.get('k_value')

        n = float(n_val.replace(',', '.')) if n_val else 1.0
        s = float(s_val.replace(',', '.')) if s_val else 1.0
        k = float(k_val.replace(',', '.')) if k_val else 1.0

        if n_val: input_values['n_value'] = n
        if s_val: input_values['s_value'] = s
        if k_val: input_values['k_value'] = k

<<<<<<< HEAD
        score = calculate_score(indicator_type, base_weight, n, s, k)
=======
        if indicator_type in ('fixed_value', 'fixed'):
            score = base_weight
        elif indicator_type == 'simple_multiplication':
            score = base_weight * n
        elif indicator_type in ('percentage_update', 'positive_feedback'):
            percent = s if s <= 1.0 else s / 100.0
            score = base_weight * n * percent
        elif indicator_type == 'scientific_publication':
            score = base_weight * n * k * s
        elif indicator_type == 'quantity_share':
            score = base_weight * n * s
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620

    except Exception as e:
        flash(f"Помилка розрахунку: {e}", 'error')
        return redirect(url_for('index'))

    entries = session.get('entries', [])
    new_entry = {
        'id': indicator_id,
        'name': indicator_details.get('text', 'Невідомо'),
        'coeff': base_weight,
        'score': score,
        'block': indicator_details.get('block', 1),
        'type': indicator_type,
        'comment': comment,
        **input_values
    }
    entries.append(new_entry)
    session['entries'] = entries
    session.modified = True
    flash(f"Додано: {score:.2f} балів", 'success')
    return redirect(url_for('index'))

@app.route('/delete/<int:entry_index>', methods=['POST'])
def delete_entry(entry_index):
    entries = session.get('entries', [])
    if 0 <= entry_index < len(entries): 
        entries.pop(entry_index)
        session['entries'] = entries
        session.modified = True
    return redirect(url_for('index'))

@app.route('/edit/<int:entry_index>', methods=['GET'])
def edit_entry(entry_index):
    entries = session.get('entries', [])
    if 0 <= entry_index < len(entries):
        return render_template('edit_entry.html', entry=entries[entry_index], entry_index=entry_index)
    return redirect(url_for('index'))


@app.route('/update/<int:entry_index>', methods=['POST'])
def update_entry(entry_index):
    entries = session.get('entries', [])
    if not (0 <= entry_index < len(entries)): return redirect(url_for('index'))

    entry_to_update = entries[entry_index]
    indicator_details = get_indicator_details(entry_to_update['id'])

    indicator_type = indicator_details.get('formula_type', 'fixed_value')
    base_weight = indicator_details.get('weight', 0)

    score = 0
    input_values = {}
    comment = request.form.get('comment', '').strip()

    try:
        n_val = request.form.get('n_value')
        s_val = request.form.get('s_value')
        k_val = request.form.get('k_value')

        n = float(n_val.replace(',', '.')) if n_val else 1.0
        s = float(s_val.replace(',', '.')) if s_val else 1.0
        k = float(k_val.replace(',', '.')) if k_val else 1.0

        if n_val: input_values['n_value'] = n
        if s_val: input_values['s_value'] = s
        if k_val: input_values['k_value'] = k

<<<<<<< HEAD
        score = calculate_score(indicator_type, base_weight, n, s, k)
=======
        if indicator_type in ('fixed_value', 'fixed'):
            score = base_weight
        elif indicator_type == 'simple_multiplication':
            score = base_weight * n
        elif indicator_type in ('percentage_update', 'positive_feedback'):
            percent = s if s <= 1.0 else s / 100.0
            score = base_weight * n * percent
        elif indicator_type == 'scientific_publication':
            score = base_weight * n * k * s
        elif indicator_type == 'quantity_share':
            score = base_weight * n * s
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620

    except Exception as e:
        flash(f"Помилка оновлення: {e}", 'error')
        return redirect(url_for('index'))

    entries[entry_index].update({
        'name': indicator_details.get('text', 'Невідомо'),
        'coeff': base_weight,
        'type': indicator_type,
        'score': score,
        'comment': comment,
        **input_values
    })
    session['entries'] = entries
    session.modified = True
    return redirect(url_for('index'))

@app.route('/upload_csv', methods=['POST'])
def upload_csv():
    if 'csv_file' not in request.files: return redirect(url_for('index'))
    file = request.files['csv_file']
    if file.filename == '' or not allowed_file(file.filename): return redirect(url_for('index'))
    
    try:
        stream = io.StringIO(file.stream.read().decode("utf-8-sig"), newline=None)
        csv_reader = csv.reader(stream, delimiter=';')
        
        parsed_entries = []
        personal_info = {}
        header_found = False
        
        for row in csv_reader:
            if not row: continue
            if "ПІБ:" in row[0]: personal_info['full_name'] = row[1]
            if "Посада:" in row[0]: personal_info['position'] = row[1]
            if "Блок" in row[0] and "Показник" in row[2]: header_found = True; continue
            
            if header_found and len(row) >= 6:
                if "Підсумки" in row[2]: break
                try:
                    entry_id = row[1].strip()
                    if not entry_id: continue
                    details = get_indicator_details(entry_id)
                    if not details: continue
                    
                    score = float(row[5].replace(',', '.'))
                    comment = row[6] if len(row) > 6 else ""
                    
                    parsed_inputs = parse_input_data_string(row[3], details['type'])
                    
                    parsed_entries.append({
                        'id': entry_id, 'name': details['name'], 'coeff': details['coeff'],
                        'score': score, 'block': details['block'], 'type': details['type'],
                        'comment': comment, **parsed_inputs
                    })
                except: pass
        
        session['entries'] = parsed_entries
        session.update(personal_info)
        session.modified = True
        flash(f"Завантажено {len(parsed_entries)} записів.", 'success')
    except Exception as e:
<<<<<<< HEAD
        logging.error(f"Error processing CSV upload: {e}")
=======
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620
        flash(f"Помилка CSV: {e}", 'error')
        
    return redirect(url_for('index'))

@app.route('/table')
def show_table():
    entries = session.get('entries', [])
    full_name = session.get('full_name')
    position = session.get('position')
    
    total = sum(e['score'] for e in entries)
    
    if full_name and position:
        leaderboard = load_leaderboard()
<<<<<<< HEAD
        # ALWAYS Update leaderboard with the current score (allows reducing score)
        leaderboard[full_name] = {'score': total, 'position': position}
        save_leaderboard(leaderboard)
        
        # --- ЗБЕРЕЖЕННЯ JSON ДЛЯ WORD ---
        try:
            safe_name = "".join([c for c in full_name if c.isalnum() or c in ' .-_']).strip()
            with open(f"details_{safe_name}.json", 'w', encoding='utf-8') as f:
                json.dump(entries, f, ensure_ascii=False, indent=4)
        except Exception as e: 
            logging.error(f"JSON save error: {e}")
        # --------------------------------
=======
        if full_name not in leaderboard or total > leaderboard[full_name]['score']:
            leaderboard[full_name] = {'score': total, 'position': position}
            save_leaderboard(leaderboard)
            
            # --- ЗБЕРЕЖЕННЯ JSON ДЛЯ WORD ---
            try:
                safe_name = "".join([c for c in full_name if c.isalnum() or c in ' .-_']).strip()
                with open(f"details_{safe_name}.json", 'w', encoding='utf-8') as f:
                    json.dump(entries, f, ensure_ascii=False, indent=4)
            except Exception as e: print(f"JSON save error: {e}")
            # --------------------------------
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620

    total_block1 = sum(e['score'] for e in entries if e.get('block') == 1)
    total_block2 = sum(e['score'] for e in entries if e.get('block') == 2)
    
    return render_template('results_table.html', 
                           entries=entries, 
                           grand_total=total, 
                           total_block1=total_block1,
                           total_block2=total_block2,
                           personal_info={'full_name': full_name, 'position': position})

@app.route('/clear')
def clear_entries(): 
    session.clear()
    return redirect(url_for('index'))

@app.route('/leaderboard')
def show_leaderboard():
    leaderboard = load_leaderboard()
    
    all_pos = sorted(list(set(d.get('position', '') for d in leaderboard.values()) - {''}))
    filter_pos = request.args.get('position_filter')
    
    lb_list = []
    for name, data in leaderboard.items():
        if filter_pos and data.get('position') != filter_pos: continue
        lb_list.append({'name': name, 'score': data.get('score', 0), 'position': data.get('position')})
    
    lb_list.sort(key=lambda x: x['score'], reverse=True)
    return render_template('leaderboard.html', leaderboard=lb_list, available_positions=all_pos, current_filter=filter_pos)

@app.route('/delete_leaderboard_entry', methods=['POST'])
def delete_leaderboard_entry():
    name = request.form.get('name')
    password = request.form.get('admin_password')
    
    if password != ADMIN_PASSWORD:
        flash('Невірний пароль!', 'error')
        return redirect(url_for('show_leaderboard'))
        
    lb = load_leaderboard()
    if name in lb:
        del lb[name]
        save_leaderboard(lb)
        flash(f'Користувача {name} видалено.', 'success')
    return redirect(url_for('show_leaderboard'))

@app.route('/download_report_docx/<name>')
def download_report_docx(name):
    # Очищуємо ім'я файлу
    safe_name = "".join([c for c in name if c.isalnum() or c in ' .-_']).strip()
    filename = f"details_{safe_name}.json"
    
    # Перевірка наявності файлу
    if not os.path.exists(filename):
        flash("Детальний звіт не знайдено. Натисніть 'Фінальна таблиця' для оновлення.", 'error')
        return redirect(url_for('show_leaderboard'))
        
    try:
        # ШВИДКЕ читання
        with open(filename, 'r', encoding='utf-8') as f:
            entries = json.load(f)
            
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # --- ШАПКА ---
        p = doc.add_paragraph('УНІВЕРСИТЕТ ЕКОНОМІКИ І ПІДПРИЄМНИЦТВА')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = doc.add_paragraph('ДОДАТОК Б')
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        p = doc.add_paragraph('Індивідуальні дані,')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        p = doc.add_paragraph('що відображають результати діяльності науково-педагогічного працівника')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'\nПІБ викладача: {name}')
        doc.add_paragraph(f'Дата формування: {datetime.now().strftime("%d.%m.%Y")}\n')
        
        # --- ТАБЛИЦЯ ---
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        
        hdr = table.rows[0].cells
        headers = ['Блок', '№', 'Показник', 'Введені дані', 'Бал', 'Коментар']
        for i, h_text in enumerate(headers):
            hdr[i].text = h_text
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        total = 0
        
        # --- ОПТИМІЗОВАНЕ СОРТУВАННЯ ---
        def smart_sort_key(e):
            block = e.get('block', 0)
            id_str = str(e.get('id', ''))
            parts = []
            for part in id_str.split('.'):
                if part.isdigit():
                    parts.append(int(part))
                else:
                    parts.append(part)
            return (block, str(parts)) 
            
        sorted_entries = sorted(entries, key=smart_sort_key)

        # Заповнення таблиці
        for e in sorted_entries:
            row_cells = table.add_row().cells
            row_cells[0].text = str(e.get('block', ''))
            row_cells[1].text = str(e.get('id', ''))
            row_cells[2].text = str(e.get('name', ''))
            
            parts = []
            if e.get('n_value') is not None: parts.append(f"n={e['n_value']}")
            if e.get('s_value') is not None: parts.append(f"S={e['s_value']}")
            if e.get('k_value') is not None: parts.append(f"k={e['k_value']}")
            if e.get('boolean_value'): parts.append("Так")
            
            row_cells[3].text = ", ".join(parts) if parts else "-"
            
            score = e.get('score', 0)
            row_cells[4].text = "{:.2f}".format(score)
            row_cells[5].text = e.get('comment', '')
            total += score
            
        # --- ПІДВАЛ ---
        doc.add_paragraph('\n')
        p = doc.add_paragraph(f'Загальний рейтинг: {total:.2f} балів')
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)
        
        doc.add_paragraph('\n' + '-'*60 + '\n')
        p = doc.add_paragraph('Індивідуальні дані, що відображають результати моєї науково-педагогічної діяльності внесені мною особисто та є достовірними.')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph('\n\n')
        sig_table = doc.add_table(rows=1, cols=2)
        sig_table.allow_autofit = True
        
        for row in sig_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
        
        sig_table.rows[0].cells[0].text = f"Викладач _______________ {name}"
        sig_table.rows[0].cells[1].text = "Завідувач кафедри _______________"

        f_out = io.BytesIO()
        doc.save(f_out)
        f_out.seek(0)
        
        # --- ВИПРАВЛЕННЯ КОДУВАННЯ (URL Encode) ---
        encoded_filename = quote(f'Report_{safe_name}.docx')
        
        return Response(
            f_out,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            # Важливий рядок: filename*=UTF-8''... дозволяє передавати кирилицю
            headers={'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"}
        )

    except Exception as e:
        print(f"Word generation error: {e}")
        flash(f"Помилка при створенні Word файлу: {e}", 'error')
        return redirect(url_for('show_leaderboard'))

@app.route('/download/csv')
def download_csv():
    entries = session.get('entries', [])
    personal_info = { 
        'full_name': session.get('full_name', 'N/A'), 
        'institution_type': session.get('institution_type', 'N/A'), 
        'department': session.get('department', 'N/A'), 
        'position': session.get('position', 'N/A') 
    }
    total_block1 = sum(e['score'] for e in entries if e.get('block') == 1)
    total_block2 = sum(e['score'] for e in entries if e.get('block') == 2)
    grand_total = total_block1 + total_block2

    output = io.StringIO()
    writer = csv.writer(output, delimiter=';', quotechar='"', quoting=csv.QUOTE_MINIMAL)
<<<<<<< HEAD
    writer.writerow(['ПІБ:', escape_csv_injection(personal_info['full_name'])])
    writer.writerow(['Заклад:', escape_csv_injection(personal_info['institution_type'])])
    writer.writerow(['Кафедра:', escape_csv_injection(personal_info['department'])])
    writer.writerow(['Посада:', escape_csv_injection(personal_info['position'])])
=======
    writer.writerow(['ПІБ:', personal_info['full_name']])
    writer.writerow(['Заклад:', personal_info['institution_type']])
    writer.writerow(['Кафедра:', personal_info['department']])
    writer.writerow(['Посада:', personal_info['position']])
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620
    writer.writerow([])
    
    header = ["Блок", "№ Пункту", "Показник", "Введені дані", "Коеф./База", "Отримані бали", "Коментар"]
    writer.writerow(header)

    def format_input_data(entry):
        parts = []
        if entry.get('type') == 'fixed': return "-"
        if entry.get('n_value') is not None: parts.append(f"n={str(entry.get('n_value', '')).replace('.', ',')}")
        if entry.get('s_value') is not None: parts.append(f"S={str(entry.get('s_value', '')).replace('.', ',')}")
        if entry.get('k_value') is not None: parts.append(f"k={str(entry.get('k_value', '')).replace('.', ',')}")
        if 'boolean_value' in entry: parts.append("Так" if entry.get('boolean_value') else "Ні")
        return " ".join(parts) if parts else "-"

    for entry in entries:
        input_data_str = format_input_data(entry)
        coeff_str = str(entry.get('coeff', '')).replace('.', ',')
        score_str = "{:.2f}".format(entry.get('score', 0)).replace('.', ',')
<<<<<<< HEAD
        comment = escape_csv_injection(entry.get('comment', ''))
        name = escape_csv_injection(entry.get('name', ''))
        row = [entry.get('block', ''), entry.get('id', ''), name, escape_csv_injection(input_data_str), coeff_str, score_str, comment]
=======
        comment = entry.get('comment', '')
        row = [entry.get('block', ''), entry.get('id', ''), entry.get('name', ''), input_data_str, coeff_str, score_str, comment]
>>>>>>> f69cf2f3143ed2211846546383e42fd0b4950620
        writer.writerow(row)
        
    writer.writerow([])
    writer.writerow(["", "", "--- Підсумки ---", "", "", "", ""])
    writer.writerow(["", "", "Всього за Блок 1:", "", "", "{:.2f}".format(total_block1).replace('.', ','), ""])
    writer.writerow(["", "", "Всього за Блок 2:", "", "", "{:.2f}".format(total_block2).replace('.', ','), ""])
    writer.writerow(["", "", "Загальна сума балів:", "", "", "{:.2f}".format(grand_total).replace('.', ','), ""])
    
    output.seek(0)
    return Response(
        u'\ufeff'.encode('utf-8') + output.getvalue().encode('utf-8'), 
        mimetype="text/csv; charset=utf-8", 
        headers={"Content-Disposition": "attachment;filename=rating_results.csv"}
    )

@app.route('/download/html')
def download_html():
    entries = session.get('entries', [])
    personal_info = {
        'full_name': session.get('full_name'),
        'institution_type': session.get('institution_type'),
        'department': session.get('department'),
        'position': session.get('position')
    }

    if not entries and not personal_info.get('full_name'):
        flash("Немає даних для завантаження.", 'warning')
        return redirect(url_for('show_table'))

    total_block1 = sum(e['score'] for e in entries if e.get('block') == 1)
    total_block2 = sum(e['score'] for e in entries if e.get('block') == 2)
    grand_total = total_block1 + total_block2

    def entry_sort_key(entry):
        block = entry.get('block', 0)
        item_key = entry.get('id', '')
        parts = item_key.split('.')
        key_tuple = [block]
        for part in parts:
            try:
                key_tuple.append(int(part))
            except ValueError:
                key_tuple.append(float('inf'))
                key_tuple.append(part)
        return tuple(key_tuple)
    sorted_entries = sorted(entries, key=entry_sort_key)

    try:
        html_content = render_template(
            'printable_table.html',
            entries=sorted_entries,
            total_block1=total_block1,
            total_block2=total_block2,
            grand_total=grand_total,
            personal_info=personal_info
        )
        return Response(
            html_content.encode('utf-8'),
            mimetype="text/html",
            headers={"Content-Disposition": "attachment;filename=rating_results.html"}
        )
    except Exception as e:
        print(f"HTML Error: {e}")
        flash(f"Помилка генерації HTML: {e}", "error")
        return redirect(url_for('show_table'))

if __name__ == '__main__':
    if not os.path.exists(LEADERBOARD_FILE): save_leaderboard({})
    app.run(host='0.0.0.0', port=5000)